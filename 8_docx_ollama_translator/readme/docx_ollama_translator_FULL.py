import os
import json
import requests
from datetime import datetime
from tqdm import tqdm
from docx import Document


MODEL = "todorov/bggpt:9B-IT-v1.0.Q6_K" # The language model for translation can be changed by replacing this name with another model name from cmd "ollama list"
API_URL = "http://localhost:11434/api/generate" #BgGPT - INSAIT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
PROMPTS_DIR = os.path.join(BASE_DIR, "prompts")
OUTPUT_BASE_DIR = os.path.join(BASE_DIR, "translated_docx_files")
LOGS_DIR = os.path.join(BASE_DIR, "logs")
PROMPTS_JSON = os.path.join(PROMPTS_DIR, "prompt_list.json")

for folder in [INPUT_DIR, PROMPTS_DIR, OUTPUT_BASE_DIR, LOGS_DIR]:
    os.makedirs(folder, exist_ok=True)

RUN_TS = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
LOG_FILE = os.path.join(LOGS_DIR, f"log_{RUN_TS}.txt")
SESSION_OUTPUT_DIR = os.path.join(OUTPUT_BASE_DIR, RUN_TS)

def write_log(message):
    """Writes a log entry with timestamp."""
    now = datetime.now().strftime("%H:%M:%S")
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{now}] {message}\n")

def get_lang_suffix(prompt_filename):
    """
    Extracts the language suffix from filename.
    Example: 'translate_en.txt' -> 'en'
    """
    name_part = prompt_filename.replace(".txt", "")
    parts = name_part.split("_")
    return parts[-1] if len(parts) > 1 else "tr"

def load_prompt_text(filename):
    """Reads the content of a prompt file from the prompts folder."""
    path = os.path.join(PROMPTS_DIR, filename)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    return None

def call_ollama(text, prompt_template, session):
    """Sends a request to Ollama to translate a specific text element."""
    if not text.strip() or len(text.strip()) < 2:
        return text
    
    if "{text}" in prompt_template:
        full_prompt = prompt_template.replace("{text}", text)
    else:
        full_prompt = f"{prompt_template}\n\nText: {text}"
    
    try:
        payload = {"model": MODEL, "prompt": full_prompt, "stream": False}
        r = session.post(API_URL, json=payload, timeout=180)
        return r.json().get("response", "").strip()
    except Exception as e:
        write_log(f"API Error: {str(e)}")
        return text

def translate_docx(file_name, lang_item, prompt_text, session):
    """Processes a single DOCX file element by element."""
    source_path = os.path.join(INPUT_DIR, file_name)
    suffix = get_lang_suffix(lang_item['name'])
    
    name_part, ext = os.path.splitext(file_name)
    new_filename = f"{name_part}_{suffix}{ext}"
    save_path = os.path.join(SESSION_OUTPUT_DIR, new_filename)
    
    write_log(f"Starting translation: {file_name} -> {new_filename}")
    
    try:
        doc = Document(source_path)
        
        elements = []
        for p in doc.paragraphs:
            if p.text.strip(): elements.append(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): elements.append(cell)

        pbar = tqdm(total=len(elements), desc=f"  ↳ {file_name[:25]}", unit="el", leave=False, colour='cyan')

        for obj in elements:
            obj.text = call_ollama(obj.text, prompt_text, session)
            pbar.update(1)
        
        doc.save(save_path)
        pbar.close()
        write_log(f"Successfully completed: {new_filename}")
        return True
    except Exception as e:
        write_log(f"Critical error processing {file_name}: {str(e)}")
        return False

def main():
    write_log("--- SESSION START ---")
    
    # 1. Load JSON list of prompts
    if not os.path.exists(PROMPTS_JSON):
        print(f"❌ Error: {PROMPTS_JSON} not found!")
        write_log("Error: prompt_list.json missing")
        return

    with open(PROMPTS_JSON, "r", encoding="utf-8") as f:
        languages = json.load(f)

    # 2. Check for DOCX files
    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx")]
    if not files:
        print(f"⚠️ No .docx files found in {INPUT_DIR}")
        write_log("End: No files to process.")
        return

    # 3. User interface
    print(f"\n📂 Files found for translation: {len(files)}")
    print("🌍 Select target language:")
    for idx, lang in enumerate(languages, 1):
        print(f"{idx:2}. {lang['desc']}")

    choice = input("\nEnter number(s) (e.g., 1,3) or 'all': ").strip().lower()
    
    if choice == 'all':
        selected_langs = languages
    else:
        try:
            indices = [int(i.strip()) - 1 for i in choice.split(",")]
            selected_langs = [languages[i] for i in indices if 0 <= i < len(languages)]
        except:
            print("❌ Invalid selection.")
            return

    if not selected_langs: return

    os.makedirs(SESSION_OUTPUT_DIR, exist_ok=True)
    write_log(f"Selected languages: {[l['name'] for l in selected_langs]}")

    # 4. Translation process
    session = requests.Session()
    print(f"\n🚀 Starting... Results will be in: {RUN_TS}")

    for lang in tqdm(selected_langs, desc="OVERALL PROGRESS (Languages)", unit="lang"):
        prompt_content = load_prompt_text(lang['name'])
        
        if not prompt_content:
            write_log(f"Error: Missing .txt prompt file for {lang['name']}")
            continue
            
        for doc_file in files:
            translate_docx(doc_file, lang, prompt_content, session)

    write_log("--- SESSION END ---")
    print("-" * 30)
    print(f"\n✨ Done! Check logs in the 'logs' folder for details.")
    print(f"📂 Results saved in folder: docx_translated_files")
    print("-" * 30)
    print("👋 Byee!")

if __name__ == "__main__":
    main()
