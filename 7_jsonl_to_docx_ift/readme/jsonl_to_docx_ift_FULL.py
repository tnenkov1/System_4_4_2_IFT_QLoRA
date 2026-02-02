import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from tqdm import tqdm

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "jsonl_files")
OUTPUT_BASE_DIR = os.path.join(BASE_DIR, "docx_files")

def apply_115_spacing(paragraph):
    """Apply 1.15 line spacing with no extra spacing before/after"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_heading_115(doc, text):
    h = doc.add_heading(text, level=1)
    apply_115_spacing(h)
    return h


def add_paragraph_115(doc, text):
    p = doc.add_paragraph(text)
    apply_115_spacing(p)
    return p

def convert_jsonl_to_docx(input_path, output_path):
    """Convert JSONL (instruction/input/output) to DOCX with 1.15 spacing only"""
    if not os.path.exists(input_path):
        return 0

    doc = Document()
    examples_count = 0

    with open(input_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue

            try:
                data = json.loads(line)
                instruction = data.get("instruction", "").strip()
                input_text = data.get("input", "").strip()
                output_text = data.get("output", "").strip()

                if instruction:
                    add_heading_115(doc, instruction)

                if input_text:
                    add_paragraph_115(doc, f"// {input_text}")

                if output_text:
                    add_paragraph_115(doc, output_text)

                add_paragraph_115(doc, "")

                examples_count += 1

            except json.JSONDecodeError:
                continue

    if examples_count > 0:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

    return examples_count

def main():
    if not os.path.exists(INPUT_DIR):
        print(f"⚠️ Input directory '{INPUT_DIR}' does not exist.")
        return

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_dir = os.path.join(OUTPUT_BASE_DIR, timestamp)

    jsonl_files = []
    for root, _, files in os.walk(INPUT_DIR):
        for f in files:
            if f.lower().endswith(".jsonl"):
                jsonl_files.append(os.path.join(root, f))

    if not jsonl_files:
        print("⚠️ No JSONL files found.")
        return

    print("\n" + "-" * 30)
    print(f"🔄 Converting {len(jsonl_files)} file(s)...")

    total_converted = 0

    for file_path in tqdm(jsonl_files, desc="Converting", unit="file"):
        filename = os.path.basename(file_path)
        out_p = os.path.join(
            output_dir,
            filename.replace(".jsonl", ".docx")
        )

        try:
            total_converted += convert_jsonl_to_docx(file_path, out_p)
        except Exception as e:
            tqdm.write(f"⚠️ Error processing '{filename}': {e}")

    print("-" * 30)
    print(f"📝 Total restored examples:: {total_converted}")
    print(f"📂 Results saved in folder: docx_files")
    print("-" * 30)
    print("👋 Byee!")


if __name__ == "__main__":
    main()
