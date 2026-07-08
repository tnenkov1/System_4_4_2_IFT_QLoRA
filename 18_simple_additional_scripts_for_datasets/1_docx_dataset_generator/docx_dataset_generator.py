import datetime
import logging
import os
import re
from time import sleep
from docx import Document
import requests
from tqdm import tqdm

# ==========================================
# CONFIGURATION
# ==========================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
OUTPUT_DIR = os.path.join(BASE_DIR, "expanded_docx")

OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_TAGS_URL = "http://localhost:11434/api/tags"

# Create folders if they do not exist
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("process.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)

# ==========================================
# OLLAMA
# ==========================================

def get_models():
    try:
        r = requests.get(OLLAMA_TAGS_URL, timeout=5)
        r.raise_for_status()
        return [m["name"] for m in r.json().get("models", [])]
    except Exception as e:
        logging.error(f"Failed to connect to Ollama: {e}")
        return []


def query_ollama(model, prompt, retries=2):
    for i in range(retries + 1):
        try:
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.3, "num_predict": 1200},
            }
            r = requests.post(OLLAMA_API_URL, json=payload, timeout=400)
            r.raise_for_status()
            return r.json()["response"]
        except Exception as e:
            if i < retries:
                logging.warning(
                    "Request error, retrying in 5 seconds..."
                )
                sleep(5)
            else:
                raise e

# ==========================================
# DOCX LOGIC
# ==========================================

def extract_sections(doc):
    sections = []
    current_title = None
    current_text = []

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            if current_title:
                sections.append((current_title, "\n".join(current_text)))
            current_title = p.text.strip()
            current_text = []
        else:
            if p.text.strip():
                current_text.append(p.text.strip())

    if current_title:
        sections.append((current_title, "\n".join(current_text)))

    return sections

def expand_section_with_qa(model, title, text):
    prompt = f"""
You are an expert analyst and trainer.

TASK:
Expand the content under the given heading by adding
EXACTLY 3 questions and EXACTLY 3 detailed answers.

RULES:
1. Each question must deepen or clarify the topic.
2. Each answer must be between 3 and 5 sentences long.
3. DO NOT repeat the original text verbatim.
4. DO NOT use lists, bullet points, or Markdown formatting.
5. Write exclusively in English.

HEADING:
{title}

ORIGINAL TEXT:
{text}

OUTPUT FORMAT (STRICT):
Question 1: ...
Answer: ...

Question 2: ...
Answer: ...

Question 3: ...
Answer: ...
"""
    return query_ollama(model, prompt)

def clean_qa_line(line: str) -> str:
    line = line.strip()
    line = re.sub(r"^Question\s*\d*\s*:\s*", "", line, flags=re.IGNORECASE)
    line = re.sub(r"^Answer\s*:\s*", "", line, flags=re.IGNORECASE)
    return line

# ==========================================
# MAIN PROCESSING
# ==========================================

def process_docx(input_path, output_path, model):
    try:
        doc = Document(input_path)
        sections = extract_sections(doc)

        if not sections:
            logging.warning(f"No headings found in: {input_path}")
            return False

        final_doc = Document()
        final_doc.add_heading(
            f"Expanded Document: {os.path.basename(input_path)}", level=0
        )

        for title, text in sections:
            # Original heading
            final_doc.add_heading(title, level=1)

            # Original text
            if text:
                final_doc.add_paragraph(text)

            # Expansion via Q&A
            qa_text = expand_section_with_qa(model, title, text)

            current_question = None
            current_answer = []

            for raw_line in qa_text.split("\n"):
                line = raw_line.strip()
                if not line:
                    continue

                if line.lower().startswith("question"):
                    if current_question and current_answer:
                        final_doc.add_heading(current_question, level=2)
                        final_doc.add_paragraph(" ".join(current_answer))

                    current_question = clean_qa_line(line)
                    current_answer = []

                elif line.lower().startswith("answer"):
                    cleaned = clean_qa_line(line)
                    if cleaned:
                        current_answer.append(cleaned)

                else:
                    current_answer.append(line)

            if current_question and current_answer:
                final_doc.add_heading(current_question, level=2)
                final_doc.add_paragraph(" ".join(current_answer))

        final_doc.save(output_path)
        return True

    except Exception as e:
        logging.error(f"Error processing {input_path}: {e}")
        return False

# ==========================================
# MAIN
# ==========================================

def main():
    models = get_models()
    if not models:
        print("❌ No models available in Ollama.")
        return

    print("\nAvailable models:")
    for i, m in enumerate(models, 1):
        print(f"{i}. {m}")

    try:
        model = models[int(input("\nSelect a model: ")) - 1]
    except Exception:
        print("Invalid choice.")
        return

    files = [f for f in os.listdir(INPUT_DIR) if f.endswith(".docx")]
    if not files:
        print("❌ No DOCX files found for processing.")
        return

    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
    out_dir = os.path.join(OUTPUT_DIR, timestamp)
    os.makedirs(out_dir, exist_ok=True)

    # ====== Counting sections ======
    total_sections = 0
    file_sections_map = {}
    for f in files:
        input_p = os.path.join(INPUT_DIR, f)
        doc = Document(input_p)
        sections = extract_sections(doc)
        file_sections_map[f] = sections
        total_sections += len(sections)

    print(
        f"\n🚀 Processing {len(files)} file(s) ({total_sections} sections in total)...\n"
    )
    success = 0

    # ====== Single progress bar ======
    with tqdm(total=total_sections, desc="Processing sections") as pbar:
        for f in files:
            input_p = os.path.join(INPUT_DIR, f)
            out_p = os.path.join(out_dir, f"Expanded_{f}")
            sections = file_sections_map[f]

            try:
                final_doc = Document()
                final_doc.add_heading(
                    f"Expanded Document: {os.path.basename(f)}", level=0
                )
            except Exception as e:
                logging.error(f"Failed to create DOCX for {f}: {e}")
                continue

            for title, text in sections:
                # Original heading
                final_doc.add_heading(title, level=1)
                if text:
                    final_doc.add_paragraph(text)

                # Q&A expansion with fallback
                try:
                    qa_text = expand_section_with_qa(model, title, text)
                except Exception as e:
                    logging.error(
                        f"Ollama did not respond for section '{title}' in {f}: {e}"
                    )
                    qa_text = "⚠️ Failed to process this section."

                current_question = None
                current_answer = []
                for raw_line in qa_text.split("\n"):
                    line = raw_line.strip()
                    if not line:
                        continue

                    if line.lower().startswith("question"):
                        if current_question and current_answer:
                            final_doc.add_heading(current_question, level=2)
                            final_doc.add_paragraph(" ".join(current_answer))
                        current_question = clean_qa_line(line)
                        current_answer = []
                    elif line.lower().startswith("answer"):
                        cleaned = clean_qa_line(line)
                        if cleaned:
                            current_answer.append(cleaned)
                    else:
                        current_answer.append(line)

                if current_question and current_answer:
                    final_doc.add_heading(current_question, level=2)
                    final_doc.add_paragraph(" ".join(current_answer))

                # ✅ Update progress for each section
                pbar.update(1)

            try:
                final_doc.save(out_p)
                success += 1
            except Exception as e:
                logging.error(f"Failed to save {out_p}: {e}")

    print(f"\n🎉 Done: {success}/{len(files)} file(s) processed successfully.")
    print(f"📁 Results saved in: {out_dir}")


if __name__ == "__main__":
    main()