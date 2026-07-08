import os
from docx import Document

# ==========================================
# FOLDERS
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
OUTPUT_DIR = os.path.join(BASE_DIR, "docx_cleaned")

# Create folders if they do not exist
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ==========================================
# CLEANING FUNCTION
# ==========================================
def remove_heading1(doc):
    """Removes all Heading 1 sections and the text beneath them until the next Heading."""
    new_doc = Document()
    skip = False

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading 1"):
            skip = True  # Start skipping content
            continue

        # If we reach another Heading type (H2, H3, etc.), stop skipping
        if (
            skip
            and p.style.name.startswith("Heading")
            and not p.style.name.startswith("Heading 1")
        ):
            skip = False

        if not skip:
            new_doc.add_paragraph(p.text, style=p.style)

    return new_doc


# ==========================================
# FILE PROCESSING
# ==========================================
def process_files():
    files = [f for f in os.listdir(INPUT_DIR) if f.endswith(".docx")]

    if not files:
        print(f"❌ No DOCX files found in folder '{INPUT_DIR}'.")
        return

    for f in files:
        doc = Document(os.path.join(INPUT_DIR, f))
        cleaned_doc = remove_heading1(doc)
        cleaned_doc.save(os.path.join(OUTPUT_DIR, f"cleaned_{f}"))

    print(
        f"✅ Done! Processed files: {len(files)}. Results are saved in '{OUTPUT_DIR}'."
    )


# ==========================================
# MAIN
# ==========================================
if __name__ == "__main__":
    process_files()