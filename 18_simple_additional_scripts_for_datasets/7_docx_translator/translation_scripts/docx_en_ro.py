import os
import requests
import logging
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from time import sleep
from tqdm import tqdm
from docx import Document
from docx.text.run import Run

# ================== CONFIG ==================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
OUTPUT_BASE_DIR = os.path.join(BASE_DIR, "translated_docx_files")
LOG_DIR = os.path.join(BASE_DIR, "logs")
PROMPT_FILE = os.path.join(BASE_DIR, "prompts", "translate_en_ro.txt")  # <-- EN→RO

MODEL = ""
API_URL = "http://localhost:11434/api/generate"

MAX_WORKERS = 1
MAX_RETRIES = 3
SLEEP_BETWEEN_REQUESTS = 0.05

# ================== SETUP ==================
for d in [INPUT_DIR, LOG_DIR, os.path.dirname(PROMPT_FILE)]:
    os.makedirs(d, exist_ok=True)

RUN_TS = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
OUTPUT_DIR = os.path.join(OUTPUT_BASE_DIR, RUN_TS)
os.makedirs(OUTPUT_DIR, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(LOG_DIR, f"run_{RUN_TS}.log"), encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ================== LOAD PROMPT ==================
def load_prompt_template():
    if not os.path.exists(PROMPT_FILE):
        default_prompt = "Translate the following text from English to Romanian. Retain professional tone and formatting (bold, italic, underline, headings). Text: {TEXT}"
        with open(PROMPT_FILE, "w", encoding="utf-8") as f:
            f.write(default_prompt)
        return default_prompt
    with open(PROMPT_FILE, "r", encoding="utf-8") as f:
        return f.read()

PROMPT_TEMPLATE = load_prompt_template()

# ================== TRANSLATE ==================
def translate_text(text: str, session: requests.Session) -> str:
    if not text.strip():
        return text

    prompt = PROMPT_TEMPLATE.replace("{TEXT}", text)

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = session.post(API_URL, json={"model": MODEL, "prompt": prompt, "stream": False}, timeout=180)
            resp.raise_for_status()
            return resp.json().get("response", "").strip()
        except Exception as e:
            logger.warning(f"Attempt {attempt} failed: {e}")
            sleep(2)
    logger.error("Max retries reached. Returning original text.")
    return text

# ================== HELPER: COPY RUN FORMATTING ==================
def copy_run_format(src_run: Run, dst_run: Run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    dst_run.font.color.rgb = src_run.font.color.rgb

# ================== PROCESS DOCX ==================
def translate_docx(file_path, session: requests.Session):
    doc = Document(file_path)
    file_name = os.path.basename(file_path)
    out_path = os.path.join(OUTPUT_DIR, f"{os.path.splitext(file_name)[0]}_ro.docx")
    tmp_path = out_path + ".tmp"

    # Collect paragraphs and table cells
    targets = []
    for para in doc.paragraphs:
        if para.text.strip():
            targets.append(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    targets.append(cell)

    logger.info(f"Translating {len(targets)} elements in {file_name}...")

    # ThreadPool translation
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_obj = {executor.submit(translate_text, obj.text, session): obj for obj in targets}

        for future in tqdm(as_completed(future_to_obj), total=len(future_to_obj), desc=f"Translating {file_name[:20]}"):
            obj = future_to_obj[future]
            try:
                translated_text = future.result()
                # Preserve original formatting
                if hasattr(obj, "paragraph_format"):  # Paragraph
                    runs = list(obj.runs)
                    for r in runs:
                        r.clear()
                    run = obj.add_run(translated_text)
                    if runs:
                        copy_run_format(runs[0], run)
                else:  # Table Cell
                    cell = obj
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(translated_text)
                # Save in real-time
                doc.save(tmp_path)
                os.replace(tmp_path, out_path)
            except Exception as e:
                logger.error(f"Failed to translate element: {e}")

    logger.info(f"Saved translated file: {out_path}")

# ================== MAIN ==================
def main():
    session = requests.Session()
    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx")]

    if not files:
        logger.warning("No .docx files found in input directory.")
        return

    for fname in files:
        translate_docx(os.path.join(INPUT_DIR, fname), session)

if __name__ == "__main__":
    main()
