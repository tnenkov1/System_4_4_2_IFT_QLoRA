import datetime
import json
import os
import re
from docx import Document
import requests
from tqdm import tqdm

# ================== CONFIGURATION ==================
INPUT_DIR = "docx_files"
OUTPUT_BASE_DIR = "translated_docx_files"

SCRIPTS_DIR = "translation_scripts"
SCRIPTS_LIST_FILE = os.path.join(SCRIPTS_DIR, "scripts_list.json")
PROMPTS_DIR = "prompts"  # Folder containing your custom text prompts

API_URL = "http://localhost:11434/api/generate"
TAGS_URL = "http://localhost:11434/api/tags"

# Setup output folders
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(PROMPTS_DIR, exist_ok=True)  # Ensures prompts directory exists
RUN_TS = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
OUTPUT_DIR = os.path.join(OUTPUT_BASE_DIR, RUN_TS)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ================== OLLAMA MODELS ==================
def get_models():
    """Fetches all currently available models from the local Ollama instance."""
    try:
        r = requests.get(TAGS_URL, timeout=5)
        r.raise_for_status()
        return [m["name"] for m in r.json().get("models", [])]
    except Exception as e:
        print(f"❌ Failed to connect to Ollama: {e}")
        return []

# ================== TRANSLATION CORE ==================
def get_custom_prompt(lang_item, text):
    """Loads a custom prompt from a .txt file by converting the prefix."""
    prompt_filename = lang_item["name"].replace("docx_", "translate_").replace(".py", ".txt")
    prompt_path = os.path.join(PROMPTS_DIR, prompt_filename)

    if os.path.exists(prompt_path):
        try:
            with open(prompt_path, "r", encoding="utf-8") as f:
                template = f.read()
            
            if "{lang_desc}" in template and "{text}" in template:
                return template.format(lang_desc=lang_item["desc"], text=text)
            else:
                return f"{template.strip()}\n\nTARGET LANGUAGE:\n{lang_item['desc']}\n\nTEXT:\n{text}\n\nTRANSLATION:\n"
        except Exception as e:
            print(f"⚠️ Error reading custom prompt file {prompt_filename}: {e}")

    # Fallback with explicit strict formatting enforcement
    return f"""Instruction: Act as a precise translation engine. Translate the text below into {lang_item['desc']}. 
Strict Rules: Output ONLY the raw translation. Do not include notes, introductions, markdown formatting, quotes, or explanations.

TEXT:
{text}

TRANSLATION:"""

def clean_llm_response(response, original_text):
    """Aggressively strips formatting markers, conversational bloat, and prompt leaks."""
    if not response:
        return original_text

    response = response.strip()

    if "TRANSLATION:" in response:
        parts = response.split("TRANSLATION:")
        response = parts[-1].strip()

    response = re.sub(r"(?i).*english\s*(->|→)\s*[a-zа-я\s]+.*", "", response)
    response = re.sub(
        r"(?i).*[a-zа-я]+\s*(->|→)\s*[a-zа-я\s]+.*", "", response
    )

    prompt_leak_patterns = [
        r"(?i)task:",
        r"(?i)target language:",
        r"(?i)strict rules:",
        r"(?i)text to translate:",
        r"(?i)text:",
        r"(?i)result:",
    ]
    for pattern in prompt_leak_patterns:
        response = re.sub(pattern, "", response)

    prefixes_to_remove = [
        "here is the translation:",
        "translation:",
        "translated text:",
        "тук е преводът:",
        "превод:",
        "french:",
        "english:",
    ]
    for prefix in prefixes_to_remove:
        if response.lower().startswith(prefix):
            response = response[len(prefix) :].strip()

    lines = [line.strip() for line in response.splitlines() if line.strip()]
    response = "\n".join(lines).strip()

    if response.startswith('"') and response.endswith('"'):
        response = response[1:-1].strip()
    if response.startswith("'") and response.endswith("'"):
        response = response[1:-1].strip()

    return response if response else original_text

def translate_element(text, lang_item, model, session):
    if not text.strip() or len(text.strip()) < 2:
        return text

    prompt = get_custom_prompt(lang_item, text)

    try:
        # Added system instruction and forced temperature to 0.0 for maximum consistency
        r = session.post(
            API_URL,
            json={
                "model": model, 
                "prompt": prompt, 
                "stream": False,
                "system": "You are a literal, word-for-word document translation API. You never converse, you never explain, and you never use markdown blocks. You only output pure text translations.",
                "options": {
                    "temperature": 0.0,
                    "top_p": 0.1
                }
            },
            timeout=120,
        )
        raw_response = r.json().get("response", "").strip()
        return clean_llm_response(raw_response, text)
    except:
        return text

def update_element_text(obj, new_text):
    """Safely updates word document elements without stripping native typography styles."""
    if hasattr(obj, 'runs') and obj.runs:
        # Keep formatting layout safe by overriding the master text through the initial run execution block
        obj.runs[0].text = new_text
        for run in obj.runs[1:]:
            run.text = ""
    else:
        obj.text = new_text

def run_translation_process(file_name, lang_item, model, session):
    file_path = os.path.join(INPUT_DIR, file_name)
    doc = Document(file_path)

    targets = []
    for p in doc.paragraphs:
        if p.text.strip():
            targets.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    targets.append(cell)

    file_pbar = tqdm(
        total=len(targets),
        desc=f"  ↳ File: {file_name[:20]}",
        unit="elem",
        leave=False,
        colour="green",
    )

    for obj in targets:
        translated = translate_element(obj.text, lang_item, model, session)
        update_element_text(obj, translated)
        file_pbar.update(1)

    file_pbar.close()

    lang_folder = os.path.join(OUTPUT_DIR, lang_item["name"].replace(".py", ""))
    os.makedirs(lang_folder, exist_ok=True)
    doc.save(os.path.join(lang_folder, file_name))

# ================== MAIN INTERFACE ==================
def main():
    models = get_models()
    if not models:
        print("❌ No models available in Ollama or service is offline.")
        return

    print("\n=== SELECT OLLAMA MODEL ===")
    for i, m in enumerate(models, 1):
        print(f"{i:2}. {m}")

    try:
        model_choice = int(input("\nSelect a model number: ")) - 1
        chosen_model = models[model_choice]
    except (ValueError, IndexError):
        print("❌ Invalid model choice.")
        return

    if not os.path.exists(SCRIPTS_LIST_FILE):
        print(f"❌ Error: Missing configuration file at {SCRIPTS_LIST_FILE}")
        return
    with open(SCRIPTS_LIST_FILE, "r", encoding="utf-8") as f:
        languages = json.load(f)

    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx")]
    if not files:
        print(f"⚠️ No DOCX documents discovered in folder directory '{INPUT_DIR}'")
        return

    print(f"\n✅ Found {len(files)} file(s) ready for translation processing.")
    print("=== SELECT TARGET LANGUAGES ===")
    for idx, lang in enumerate(languages, 1):
        print(f"{idx:2}. {lang['desc']}")

    selection = (
        input("\nEnter language numbers separated by commas (e.g., 1,3,5) or 'all': ")
        .strip()
        .lower()
    )

    if selection == "all":
        chosen_langs = languages
    else:
        try:
            indices = [int(x.strip()) - 1 for x in selection.split(",")]
            chosen_langs = [
                languages[i] for i in indices if 0 <= i < len(languages)
            ]
        except:
            print("❌ Invalid configuration selection choice.")
            return

    if not chosen_langs:
        return

    session = requests.Session()
    print(
        f"\n🚀 Launching translation pipeline targeting {len(files)} file(s) across {len(chosen_langs)} language profile(s) using model {chosen_model}...\n"
    )

    for lang in tqdm(chosen_langs, desc="TOTAL PROGRESS (Languages)", unit="lang"):
        for file_name in files:
            run_translation_process(file_name, lang, chosen_model, session)

    print(f"\n✨ Process completed successfully! Output saved in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()