import datetime
import os
import re
from time import sleep
from docx import Document
import requests
from tqdm import tqdm

# ==========================================
# CONFIGURATION
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
OUTPUT_DIR = os.path.join(BASE_DIR, "conditionals_docx")

os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_TAGS_URL = "http://localhost:11434/api/tags"

# ==========================================
# HELPER FUNCTIONS
# ==========================================
def safe_filename(name):
    """Replace forbidden characters in file or folder names with '_'."""
    return re.sub(r'[\\/:*?"<>|]', "_", name)

# ==========================================
# OLLAMA COMMUNICATION
# ==========================================
def get_models():
    try:
        r = requests.get(OLLAMA_TAGS_URL, timeout=5)
        r.raise_for_status()
        return [m["name"] for m in r.json().get("models", [])]
    except Exception as e:
        print(f"❌ Failed to connect to Ollama: {e}")
        return []

def query_ollama(model, prompt, retries=2):
    for i in range(retries + 1):
        try:
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.4, "num_predict": 1000},
            }
            r = requests.post(OLLAMA_API_URL, json=payload, timeout=400)
            r.raise_for_status()
            return r.json()["response"]
        except Exception as e:
            if i < retries:
                sleep(5)
            else:
                raise e

# ==========================================
# PARAGRAPH PROCESSING
# ==========================================
def process_paragraph_text(text, model):
    """Generates unformatted conditional sentences based on a single paragraph text."""
    prompt = f"""
You are an expert.

TASK:
Read the following text and create exactly 4 original conditional sentences in English based on its meaning and themes.

STRICT RULES:
1. Include exactly one sentence using the zero conditional structure.
2. Include exactly one sentence using the first conditional structure.
3. Include exactly one sentence using the second conditional structure.
4. Include exactly one sentence using a mixed conditional structure.
5. Do not label, identify, or separate the conditional types in any way.
6. Do not use headings, prefixes, numbering, bullet points, lists, or Markdown formatting.
7. Output only the 4 sentences and nothing else.
8. Each sentence must be grammatically correct and naturally connected to the original text.
9. Avoid copying sentences or phrases from the original text; create new sentences inspired by its ideas.
10. Use standard English grammar and punctuation.
11. Place each sentence on a separate line.
12. Do not add explanations, notes, or any additional commentary.
13. Before responding, verify that all four required conditional structures are present exactly once and that no extra text is included.

ORIGINAL TEXT:
{text}
"""
    try:
        response = query_ollama(model, prompt)
        return response.strip()
    except Exception as e:
        print(f"❌ Ollama did not respond for paragraph segment: {e}")
        return "Failed to generate conditionals for this paragraph."

# ==========================================
# DOCX PROCESSING
# ==========================================
def process_docx(input_path, output_path, model):
    try:
        doc = Document(input_path)
        final_doc = Document()
        final_doc.add_heading(f"Conditionals Grammar Report: {os.path.basename(input_path)}", level=0)

        # Преброяваме параграфите за прогрес бара
        total_paragraphs = len(doc.paragraphs)

        with tqdm(doc.paragraphs, total=total_paragraphs, desc=f"Processing {os.path.basename(input_path)}", leave=False) as pbar:
            for p in pbar:
                text_clean = p.text.strip()
                if not text_clean:
                    continue

                # Ако параграфът е заглавие (Heading 1, Heading 2 и т.н.), го пренасяме директно
                if p.style.name.startswith("Heading"):
                    # Опит за извличане на нивото на заглавието (напр. "Heading 2" -> 2)
                    try:
                        level = int(p.style.name.replace("Heading ", ""))
                    except ValueError:
                        level = 1
                    
                    final_doc.add_heading(text_clean, level=min(level, 9))
                
                # Ако е обикновен текст и е достатъчно дълъг за смислен анализ
                elif len(text_clean) > 20:
                    # Генерираме условни изречения на база текста
                    llm_output = process_paragraph_text(text_clean, model)
                    
                    # Записваме резултата в новия файл под съответното заглавие
                    final_doc.add_paragraph(llm_output)
                    final_doc.add_paragraph("")  # Празен ред за разстояние

        final_doc.save(output_path)
        return True
    except Exception as e:
        print(f"❌ Error processing file {os.path.basename(input_path)}: {e}")
        return False

# ==========================================
# MAIN
# ==========================================
def main():
    models = get_models()
    if not models:
        print("❌ No models available in Ollama.")
        return

    print("\nAvailable models:")
    for i, m in enumerate(models, 1):
        print(f"{i}. {m}")

    try:
        model = models[int(input("\nSelect a model: ")) - 1]
    except Exception:
        print("Invalid choice.")
        return

    files = [f for f in os.listdir(INPUT_DIR) if f.endswith(".docx")]
    if not files:
        print(f"❌ No DOCX files found in folder '{INPUT_DIR}'.")
        return

    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
    out_dir = os.path.join(OUTPUT_DIR, timestamp)
    os.makedirs(out_dir, exist_ok=True)

    print(f"\n🚀 Processing {len(files)} file(s) with original structural layout...\n")
    success = 0

    for f in tqdm(files, desc="Files"):
        in_p = os.path.join(INPUT_DIR, f)
        out_p = os.path.join(out_dir, f"Conditionals_{safe_filename(f)}")
        if process_docx(in_p, out_p, model):
            success += 1

    print(f"\n🎉 Done: {success}/{len(files)} file(s) processed successfully.")
    print(f"📁 Results saved in: {out_dir}")

if __name__ == "__main__":
    main()