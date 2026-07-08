import datetime
import logging
import os
import re
from time import sleep
from docx import Document
import requests
from tqdm import tqdm

# ==========================================
# CONFIGURATION
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
OUTPUT_DIR = os.path.join(BASE_DIR, "structured_docx")

os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_TAGS_URL = "http://localhost:11434/api/tags"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("process.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)

# ==========================================
# HELPER FUNCTIONS
# ==========================================

def safe_filename(name):
    """Replace forbidden characters in file or folder names with '_'."""
    return re.sub(r'[\\/:*?"<>|]', "_", name)

def clean_line(line: str) -> str:
    """Cleans a line from Markdown characters (*, #, _) and trailing spaces."""
    line = line.strip()
    line = re.sub(r"^[#\*\-_ \t]+", "", line)  # Removes prefix formatting symbols
    line = re.sub(r"[*_]", "", line)  # Removes inner bold/italic markdown formatting
    return line

def is_heading(line: str) -> bool:
    """Checks if a line acts as a heading (short, no internal punctuation periods)."""
    match = re.match(r"^\d+[\.\s\-)]+\s*(.+)", line)
    if not match:
        return False
    content = match.group(1).strip()
    if len(line) > 120:
        return False
    if ". " in content[:-1]:
        return False
    return True

def split_text_to_chunks(text, max_chars=8000, overlap=500):
    """Splits text into chunks with a specified overlap."""
    chunks = []
    start = 0
    text_len = len(text)
    if text_len <= max_chars:
        return [text]
    while start < text_len:
        end = start + max_chars
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= text_len:
            break
        start = end - overlap
    return chunks

# ==========================================
# OLLAMA COMMUNICATION
# ==========================================

def get_models():
    try:
        r = requests.get(OLLAMA_TAGS_URL, timeout=5)
        r.raise_for_status()
        return [m["name"] for m in r.json().get("models", [])]
    except Exception as e:
        logging.error(f"Failed to connect to Ollama: {e}")
        return []


def query_ollama(model, text, retries=2):
    """Sends a request to the model with built-in retry logic."""
    prompt = f"""
You are a professional editor.

You must strictly follow ALL of the rules below without exception.

1. Organize the response into several main sections based on the topic.
2. Each section must have a SHORT heading (maximum 10 words).
3. Headings must use the exact format:
   "1. Heading"
   "2. Heading"
   "3. Heading"
   and so on.
4. Never use Roman numerals (I., II., III., etc.).
5. Leave exactly one blank line between each heading and the next section.
6. Under each heading, write one detailed, highly informative, and well-developed paragraph.
7. Every paragraph must begin with regular text only.
8. Paragraphs MUST NOT begin with:
   - numbers
   - digits
   - bullet points
   - dashes
   - symbols
   - letters followed by a parenthesis (e.g., A), a))
9. Do not use bullet lists, numbered lists, nested lists, tables, or subheadings inside paragraphs.
10. Do not use Markdown formatting. Forbidden elements include:
    - #
    - ##
    - ###
    - **
    - *
    - _
    - >
    - -
    - ```
11. Do not add an introduction, conclusion, notes, warnings, or any extra text unless explicitly requested.
12. Produce only the requested content.
13. If any default formatting conflicts with these instructions, these instructions always take priority.
14. Verify your output before responding. If any rule is violated, regenerate the entire response until every rule is satisfied.

TEXT TO PROCESS:
{text}
"""
    for i in range(retries + 1):
        try:
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.5, "num_predict": 3000},
            }
            r = requests.post(OLLAMA_API_URL, json=payload, timeout=500)
            r.raise_for_status()
            return r.json()["response"]
        except Exception as e:
            if i < retries:
                logging.warning(
                    f"⚠️ Request error. Attempt {i+1}/{retries}. Retrying in 5 seconds..."
                )
                sleep(5)
            else:
                raise e

# ==========================================
# DOCX PROCESSING
# ==========================================

def process_docx(input_path, output_path, model):
    try:
        doc = Document(input_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not full_text:
            logging.warning(f"File is completely empty: {input_path}")
            return False

        chunks = split_text_to_chunks(full_text)
        final_doc = Document()
        final_doc.add_heading(
            f"Structured Report: {os.path.basename(input_path)}", 0
        )
        final_doc.add_paragraph(
            f"Processing Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )

        for i, chunk in enumerate(
            tqdm(chunks, desc=f"Processing {os.path.basename(input_path)}")
        ):
            try:
                structured_response = query_ollama(model, chunk)
            except Exception as e:
                logging.error(
                    f"Ollama did not respond for a chunk of {os.path.basename(input_path)}: {e}"
                )
                structured_response = "⚠️ Failed to process this segment."

            seen_headings = set()
            for raw_line in structured_response.split("\n"):
                line = clean_line(raw_line)
                if not line:
                    continue
                if is_heading(line):
                    if line in seen_headings:
                        continue
                    final_doc.add_heading(line, level=1)
                    seen_headings.add(line)
                else:
                    final_doc.add_paragraph(line)

        final_doc.save(output_path)
        return True
    except Exception as e:
        logging.error(
            f"Critical error occurred while processing {os.path.basename(input_path)}: {e}"
        )
        return False

# ==========================================
# MAIN
# ==========================================

def main():
    models = get_models()
    if not models:
        print("❌ No connection to Ollama or no models are installed.")
        return

    print("\n--- Available Models ---")
    for i, m in enumerate(models, 1):
        print(f"{i}. {m}")

    try:
        choice = int(input("Select model number: ")) - 1
        model = models[choice]
    except (ValueError, IndexError):
        print("Invalid choice.")
        return

    files = [f for f in os.listdir(INPUT_DIR) if f.endswith(".docx")]
    if not files:
        print(f"❌ No DOCX files found in directory '{INPUT_DIR}'")
        return

    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
    session_out_dir = os.path.join(OUTPUT_DIR, timestamp)
    os.makedirs(session_out_dir, exist_ok=True)

    # ====== Aggregating all text chunks ======
    total_chunks = 0
    file_chunks_map = {}
    for f in files:
        input_p = os.path.join(INPUT_DIR, f)
        doc = Document(input_p)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        chunks = split_text_to_chunks(full_text)
        file_chunks_map[f] = chunks
        total_chunks += len(chunks)

    print(
        f"\n🚀 Starting processing on {len(files)} file(s) ({total_chunks} total chunks)...\n"
    )
    success_count = 0

    # ====== Unified progress bar ======
    with tqdm(total=total_chunks, desc="Processing all chunks") as pbar:
        for f in files:
            input_p = os.path.join(INPUT_DIR, f)
            output_p = os.path.join(
                session_out_dir, f"Structured_{safe_filename(f)}"
            )
            final_doc = Document()
            final_doc.add_heading(f"Structured Report: {os.path.basename(f)}", 0)
            final_doc.add_paragraph(
                f"Processing Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            seen_headings = set()

            for chunk in file_chunks_map[f]:
                try:
                    structured_response = query_ollama(model, chunk)
                except Exception as e:
                    logging.error(
                        f"Ollama failed to return a response for part of {f}: {e}"
                    )
                    structured_response = "⚠️ Failed to process this segment."

                for raw_line in structured_response.split("\n"):
                    line = clean_line(raw_line)
                    if not line:
                        continue
                    if is_heading(line):
                        if line in seen_headings:
                            continue
                        final_doc.add_heading(line, level=1)
                        seen_headings.add(line)
                    else:
                        final_doc.add_paragraph(line)

                pbar.update(1)  # ✅ Update progress metrics for every single chunk

            final_doc.save(output_p)
            success_count += 1

    print(
        f"\n🎉 Done! Successfully processed: {success_count}/{len(files)} file(s)"
    )
    print(f"📁 Results are located at: {session_out_dir}")

if __name__ == "__main__":
    main()