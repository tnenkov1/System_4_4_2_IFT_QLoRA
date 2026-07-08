import datetime
import logging
import os
import re
from time import sleep
from docx import Document
import requests
from tqdm import tqdm

# ==========================================
# CONFIGURATION
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "docx_files")
OUTPUT_DIR = os.path.join(BASE_DIR, "associations_docx")

os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_TAGS_URL = "http://localhost:11434/api/tags"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("process.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)

# ==========================================
# HELPER FUNCTIONS
# ==========================================
def safe_filename(name):
    """Replace forbidden characters in file or folder names with '_'."""
    return re.sub(r'[\\/:*?"<>|]', "_", name)

def extract_sections(doc):
    """Returns a list of sections (Heading, text, heading level)."""
    sections = []
    current_title = None
    current_text = []
    heading_level = 1

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            if current_title:
                sections.append(
                    (current_title, "\n".join(current_text), heading_level)
                )
            current_title = p.text.strip()
            heading_level = (
                int(p.style.name.replace("Heading ", ""))
                if p.style.name.replace("Heading ", "").isdigit()
                else 1
            )
            current_text = []
        else:
            if p.text.strip():
                current_text.append(p.text.strip())

    if current_title:
        sections.append((current_title, "\n".join(current_text), heading_level))

    return sections

# ==========================================
# OLLAMA COMMUNICATION
# ==========================================
def get_models():
    try:
        r = requests.get(OLLAMA_TAGS_URL, timeout=5)
        r.raise_for_status()
        return [m["name"] for m in r.json().get("models", [])]
    except Exception as e:
        logging.error(f"Failed to connect to Ollama: {e}")
        return []

def query_ollama(model, prompt, retries=2):
    for i in range(retries + 1):
        try:
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.5, "num_predict": 3000},
            }
            r = requests.post(OLLAMA_API_URL, json=payload, timeout=400)
            r.raise_for_status()
            return r.json()["response"]
        except Exception as e:
            if i < retries:
                logging.warning(
                    f"Request error, attempt {i+1}/{retries}. Retrying in 5 seconds..."
                )
                sleep(5)
            else:
                raise e

# ==========================================
# SECTION PROCESSING
# ==========================================
def process_section(title, text, model):
    """Generates associations for the text under a given Heading."""
    if not text.strip():
        return ["No text available for analysis."]

    prompt = f"""
You are an expert analyst and educator.

TASK:
Read the input text and generate EXACTLY 3 distinct associations or related ideas that naturally expand, explain, or connect to its main topic.

STRICT RULES:
1. Output EXACTLY 3 lines.
2. Each line must contain exactly one association.
3. Each association must be a single sentence containing no more than 15 words.
4. Write exclusively in English.
5. Do NOT use numbering, bullets, Markdown, quotation marks, labels, prefixes, or any other formatting.
6. Do NOT repeat, paraphrase, summarize, or quote the original text.
7. Each association must introduce a new but closely related idea.
8. Do not include explanations, introductions, conclusions, or extra commentary.
9. Do not leave blank lines between associations.
10. Verify that the output contains exactly three sentences before responding. If any rule is violated, regenerate the entire output.

ORIGINAL TEXT:
{text}

Start directly with the associations, outputting one sentence per line.
"""
    try:
        response = query_ollama(model, prompt)
    except Exception as e:
        logging.error(f"Ollama did not respond for section '{title}': {e}")
        response = "Failed to generate associations."

    # Filter out empty and duplicate lines
    seen = set()
    final_lines = []
    for line in response.split("\n"):
        line = line.strip()
        if line and line not in seen:
            final_lines.append(line)
            seen.add(line)

    return final_lines

# ==========================================
# DOCX PROCESSING
# ==========================================
def process_docx(input_path, output_path, model):
    try:
        doc = Document(input_path)
        final_doc = Document()
        final_doc.add_heading(
            f"Associations Report: {os.path.basename(input_path)}", level=0
        )

        sections = extract_sections(doc)

        # Progress bar tracking Headings within the file
        with tqdm(
            sections, desc=f"Processing {os.path.basename(input_path)}"
        ) as pbar:
            for title, text, level in pbar:
                final_doc.add_heading(title, level=min(level, 9))
                lines = process_section(title, text, model)
                for line in lines:
                    final_doc.add_paragraph(line)

        final_doc.save(output_path)
        return True
    except Exception as e:
        logging.error(f"Error processing file {os.path.basename(input_path)}: {e}")
        return False

# ==========================================
# MAIN
# ==========================================
def main():
    models = get_models()
    if not models:
        print("❌ No models available in Ollama.")
        return

    print("\nAvailable models:")
    for i, m in enumerate(models, 1):
        print(f"{i}. {m}")

    try:
        model = models[int(input("\nSelect a model: ")) - 1]
    except Exception:
        print("Invalid choice.")
        return

    files = [f for f in os.listdir(INPUT_DIR) if f.endswith(".docx")]
    if not files:
        print(f"❌ No DOCX files found in folder '{INPUT_DIR}'.")
        return

    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
    out_dir = os.path.join(OUTPUT_DIR, timestamp)
    os.makedirs(out_dir, exist_ok=True)

    print(f"\n🚀 Processing {len(files)} file(s)...\n")
    success = 0

    for f in tqdm(files, desc="Files"):
        in_p = os.path.join(INPUT_DIR, f)
        out_p = os.path.join(out_dir, f"Associations_{safe_filename(f)}")
        if process_docx(in_p, out_p, model):
            success += 1

    print(f"\n🎉 Done: {success}/{len(files)} file(s) processed successfully.")
    print(f"📁 Results saved in: {out_dir}")

if __name__ == "__main__":
    main()