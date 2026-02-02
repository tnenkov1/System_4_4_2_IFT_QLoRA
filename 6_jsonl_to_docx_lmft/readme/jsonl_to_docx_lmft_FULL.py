import os
import json
from datetime import datetime
from docx import Document
from tqdm import tqdm

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "jsonl_files") 
OUTPUT_BASE_DIR = os.path.join(BASE_DIR, "docx_files")

def convert_jsonl_to_docx(input_path, output_path):
    """Converts a JSONL file back into a DOCX document structure."""
    doc = Document()

    examples_count = 0

    with open(input_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue

            try:
                data = json.loads(line)
                prompt = data.get("prompt", "")
                completion = data.get("completion", "")

                if prompt:
                    # Add prompt as Heading level 1
                    doc.add_heading(prompt, level=1)

                if completion:
                    # Add completion as a regular paragraph
                    doc.add_paragraph(completion)

                examples_count += 1
            except json.JSONDecodeError:
                # Skip invalid JSON lines
                continue

    if examples_count > 0:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

    return examples_count

def main(input_dir, output_base_dir):
    if not os.path.exists(input_dir):
        print(f"⚠️ Directory {input_dir} does not exist.")
        return

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_dir = os.path.join(output_base_dir, timestamp)

    jsonl_files = []
    for root, dirs, files in os.walk(input_dir):
        for f in files:
            if f.lower().endswith(".jsonl"):
                jsonl_files.append(os.path.join(root, f))

    if not jsonl_files:
        print("⚠️ No JSONL files found for processing.")
        return

    total_converted = 0

    print("\n" + "-" * 30)
    for file_path in tqdm(jsonl_files, desc="🔄 Converting to DOCX"):
        rel_path = os.path.relpath(file_path, input_dir)
        output_filename = rel_path.replace(".jsonl", ".docx")
        out_p = os.path.join(output_dir, output_filename)

        try:
            count = convert_jsonl_to_docx(file_path, out_p)
            total_converted += count
        except Exception as e:
            tqdm.write(f"⚠️ Error processing {file_path}: {e}")

    print("-" * 30)
    print(f"📝 Total restored examples: {total_converted}")
    print(f"📂 Results saved in folder: docx_files")
    print("-" * 30)
    print("👋 Byee!")

if __name__ == "__main__":
    main(INPUT_DIR, OUTPUT_BASE_DIR)
